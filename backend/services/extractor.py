"""
services/extractor.py
---------------------
Handles safe, in-memory text extraction from uploaded files.
Supports: .txt, .csv, .pdf (pypdf), .docx (python-docx).
All files are capped at 10 MB before being touched.
"""
from __future__ import annotations

import io
import logging
from typing import Tuple

logger = logging.getLogger(__name__)

MAX_BYTES = 10 * 1024 * 1024  # 10 MB


class ExtractionError(Exception):
    """Raised when a file cannot be parsed."""


def _enforce_size(data: bytes) -> None:
    if len(data) > MAX_BYTES:
        raise ExtractionError(
            f"File too large ({len(data):,} bytes). Maximum allowed is {MAX_BYTES:,} bytes (10 MB)."
        )


def extract_text(filename: str, data: bytes) -> Tuple[str, int]:
    """
    Extract plain text from raw file bytes.

    Returns
    -------
    (text, size_bytes)
    """
    _enforce_size(data)

    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    try:
        if ext in ("txt", "csv"):
            return _extract_text_csv(data), len(data)
        elif ext == "pdf":
            return _extract_pdf(data), len(data)
        elif ext == "docx":
            return _extract_docx(data), len(data)
        else:
            # Best-effort: attempt to decode as UTF-8 text
            try:
                return data.decode("utf-8", errors="replace"), len(data)
            except Exception:
                raise ExtractionError(
                    f"Unsupported file extension '.{ext}'. Supported: .txt, .csv, .pdf, .docx"
                )
    except ExtractionError:
        raise
    except Exception as exc:
        logger.exception("Extraction failed for %s", filename)
        raise ExtractionError(f"Failed to parse file: {exc}") from exc


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _extract_text_csv(data: bytes) -> str:
    # UTF-8 with BOM / latin-1 fallback
    for enc in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError as e:  # pragma: no cover
        raise ExtractionError("pypdf is not installed") from e

    reader = PdfReader(io.BytesIO(data))
    pages: list[str] = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001
            pages.append("")
    return "\n".join(pages)


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document  # type: ignore
    except ImportError as e:  # pragma: no cover
        raise ExtractionError("python-docx is not installed") from e

    doc = Document(io.BytesIO(data))
    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)
